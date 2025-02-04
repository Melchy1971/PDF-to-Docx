import os
import pdfplumber
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, ttk


def select_pdfs():
    """Öffnet einen Dialog zur Auswahl von PDF-Dateien."""
    file_paths = filedialog.askopenfilenames(filetypes=[("PDF files", "*.pdf")])
    if file_paths:
        entry_pdf_paths.delete(0, tk.END)
        entry_pdf_paths.insert(0, ";".join(file_paths))


def select_output_dir():
    """Öffnet einen Dialog zur Auswahl eines Speicherorts für die konvertierten Dateien."""
    directory = filedialog.askdirectory()
    if directory:
        entry_output_dir.delete(0, tk.END)
        entry_output_dir.insert(0, directory)


def analyze_text_styles(page):
    """Analysiert die Schriftgrößen und formatiert den Text basierend darauf."""
    text_elements = []
    if page.chars:
        current_size = None
        current_text = []
        for char in page.chars:
            if current_size is not None and char["size"] != current_size:
                # Speichere bisher gesammelten Text mit seiner Schriftgröße
                text_elements.append({
                    "text": "".join(current_text).strip(),
                    "size": current_size
                })
                current_text = []
            current_text.append(char["text"])
            current_size = char["size"]
        
        # Füge den letzten Abschnitt hinzu
        if current_text:
            text_elements.append({
                "text": "".join(current_text).strip(),
                "size": current_size
            })

    return text_elements


def pdf_to_word(pdf_paths, output_dir):
    """Konvertiert die ausgewählten PDF-Dateien in Word-Dokumente mit Formatierung."""
    if not pdf_paths or not output_dir:
        messagebox.showerror("Fehler", "Bitte Dateien und Speicherort auswählen.")
        return

    pdf_list = pdf_paths.split(";")
    progress_bar["maximum"] = len(pdf_list)
    docx_paths = []

    try:
        for i, pdf_path in enumerate(pdf_list):
            docx_filename = os.path.join(output_dir, os.path.splitext(os.path.basename(pdf_path))[0] + ".docx")
            doc = Document()

            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    # Analysiere Text und Schriftgrößen
                    text_elements = analyze_text_styles(page)

                    for elem in text_elements:
                        text = elem["text"]
                        size = elem["size"]
                        if size > 15:  # Beispiel: große Schriftgrößen als Überschriften erster Ebene
                            doc.add_heading(text, level=1)
                        elif size > 12:  # Etwas kleinere Schrift als Überschriften zweiter Ebene
                            doc.add_heading(text, level=2)
                        else:
                            doc.add_paragraph(text)

                    # Tabellen erkennen und übernehmen
                    tables = page.extract_tables()
                    for table in tables:
                        word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                        for row_index, row in enumerate(table):
                            for col_index, cell in enumerate(row):
                                word_table.cell(row_index, col_index).text = cell

                    # Seitenumbruch nach jeder PDF-Seite
                    doc.add_page_break()

            doc.save(docx_filename)
            docx_paths.append(docx_filename)
            progress_bar["value"] = i + 1
            root.update_idletasks()

        messagebox.showinfo("Erfolg", "Die Dateien wurden erfolgreich konvertiert:\n" + "\n".join(docx_paths))

    except Exception as e:
        messagebox.showerror("Fehler", f"Beim Konvertieren ist ein Fehler aufgetreten:\n{e}")


def convert_pdfs():
    """Startet die Konvertierung."""
    pdf_paths = entry_pdf_paths.get()
    output_dir = entry_output_dir.get()
    pdf_to_word(pdf_paths, output_dir)


def exit_program():
    """Beendet das Programm."""
    root.destroy()


# GUI mit tkinter
root = tk.Tk()
root.title("PDF zu Word Konverter")
root.geometry("600x300")
root.resizable(False, False)

# PDF-Dateien auswählen
tk.Label(root, text="Wähle eine oder mehrere PDF-Dateien aus:").pack(pady=5)
entry_pdf_paths = tk.Entry(root, width=70)
entry_pdf_paths.pack(pady=5)
btn_select_pdfs = tk.Button(root, text="Durchsuchen", command=select_pdfs)
btn_select_pdfs.pack(pady=5)

# Zielverzeichnis auswählen
tk.Label(root, text="Wähle den Speicherort für die Word-Dateien:").pack(pady=5)
entry_output_dir = tk.Entry(root, width=70)
entry_output_dir.pack(pady=5)
btn_select_output_dir = tk.Button(root, text="Speicherort auswählen", command=select_output_dir)
btn_select_output_dir.pack(pady=5)

# Fortschrittsbalken
progress_bar = ttk.Progressbar(root, orient="horizontal", mode="determinate", length=400)
progress_bar.pack(pady=10)

# Konvertieren und Beenden
btn_convert = tk.Button(root, text="Wandeln", command=convert_pdfs, bg="green", fg="white")
btn_convert.pack(pady=5)
btn_exit = tk.Button(root, text="Beenden", command=exit_program, bg="red", fg="white")
btn_exit.pack(pady=5)

# Starte die GUI
root.mainloop()